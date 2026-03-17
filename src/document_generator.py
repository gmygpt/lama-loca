"""
Генератор документов: отчёты, конспекты, эссе (DOCX + Markdown)
"""
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config


class DocumentGenerator:

    def __init__(self):
        os.makedirs(config.OUTPUT_DIR, exist_ok=True)

    def _sanitize(self, name: str) -> str:
        name = re.sub(r'[^\w\s\-]', '', name)
        name = re.sub(r'\s+', '_', name.strip())
        return name[:80]

    def _ts(self) -> str:
        return datetime.now().strftime("%Y%m%d_%H%M%S")

    def save_markdown(self, text: str, topic: str, doc_type: str = "отчёт") -> str:
        filename = f"{doc_type}_{self._sanitize(topic)}_{self._ts()}.md"
        filepath = os.path.join(config.OUTPUT_DIR, filename)

        header = f"# {doc_type.capitalize()}: {topic}\n\n"
        header += f"*Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}*\n\n---\n\n"

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(header + text)
        return filepath

    def save_docx(self, text: str, topic: str, doc_type: str = "отчёт") -> str:
        filename = f"{doc_type}_{self._sanitize(topic)}_{self._ts()}.docx"
        filepath = os.path.join(config.OUTPUT_DIR, filename)

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.5

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)

        title = doc.add_heading(doc_type.capitalize(), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"Тема: {topic}")
        run.font.size = Pt(16)
        run.bold = True

        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}").font.size = Pt(12)

        doc.add_page_break()

        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+[\.\)]\s', line):
                txt = re.sub(r'^\d+[\.\)]\s', '', line)
                doc.add_paragraph(txt, style='List Number')
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                r = p.add_run(line.strip('*'))
                r.bold = True
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Cm(1.25)

        doc.save(filepath)
        return filepath

    def generate(self, text: str, topic: str, doc_type: str = "отчёт",
                 fmt: str = "both") -> list:
        files = []
        if fmt in ("md", "both"):
            files.append(self.save_markdown(text, topic, doc_type))
        if fmt in ("docx", "both"):
            files.append(self.save_docx(text, topic, doc_type))
        return files
